from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument


class ExtractionError(Exception):
    pass


def extract_text(file_path: str, file_ext: str) -> str:
    """Extract raw text from a PDF, DOCX, or TXT file."""
    ext = file_ext.lower()
    path = Path(file_path)

    if not path.exists():
        raise ExtractionError("File not found on disk.")

    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".txt":
            return _extract_txt(path)
        else:
            raise ExtractionError(f"Unsupported file extension: {ext}")
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Failed to extract text: {exc}") from exc


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    full_text = "\n".join(pages_text).strip()
    if not full_text:
        raise ExtractionError(
            "No extractable text found in PDF (it may be a scanned/image-only document)."
        )
    return full_text


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables, common in medical reports (lab result tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise ExtractionError("No extractable text found in DOCX file.")
    return full_text


def _extract_txt(path: Path) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            text = path.read_text(encoding=encoding).strip()
            if not text:
                raise ExtractionError("The TXT file is empty.")
            return text
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode TXT file with supported encodings.")
